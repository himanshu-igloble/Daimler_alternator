"""V11.1_ALT — 3-version executive summary (.docx).

Plain-language explanation + verified metrics for V10.6.2 / V11 / V11.1.
All figures match the committed caches (sources listed in the appendix table).
Output: V11.1_ALT/reports/Alternator_3Version_Summary.docx
"""
from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GOLD = RGBColor(0xC5, 0x8B, 0x1F)
OUT = pathlib.Path(__file__).resolve().parents[1] / "reports" / "Alternator_3Version_Summary.docx"


def _h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else GOLD
    return h


def _p(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def _bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(10.5)


def _table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, htxt in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = Pt(9.5)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    return t


def main() -> None:
    doc = Document()

    title = doc.add_heading("Alternator Predictive Maintenance — Three-Iteration Summary", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run("V10.6.2 (honest RUL baseline)  ·  V11 (lead-time heuristics)  ·  "
                    "V11.1 (covariate RUL test)   |   25 trucks: 10 failed, 15 in service")
    r.italic = True
    r.font.size = Pt(10)

    # ----- plain-language story -----
    _h(doc, "1. The three questions, in plain language")
    _p(doc, "For every truck we ask: WHICH truck is likely to fail? WHEN do alternators "
            "fail in general? And when will THIS specific truck fail?")
    _bullets(doc, [
        "WHICH — solved and frozen since V10.5.3: a risk score that separates failing from "
        "healthy trucks with 0.93/1.00 accuracy (AUROC). Identical in all three iterations.",
        "WHEN (fleet) — known: alternators in this fleet wear out around 601 days of service "
        "(roughly 120,000 km). Identical in all three iterations.",
        "WHEN (each truck) — the hard one. A dumb rule ('every truck fails at ~601 days, count "
        "down') predicts within ~50 days. Every sophisticated model has been off by ~140 days.",
    ])

    _h(doc, "2. What each iteration did")
    _p(doc, "V10.6.2 — proved the humbling baseline: the dumb fleet-clock rule beats the "
            "per-truck model (50 vs ~140 days error). With only 10 failures to learn from, "
            "individual failure dates were declared out of reach; the covariate idea was "
            "shelved untested.")
    _p(doc, "V11 — hunted for earlier warning signs in the same six electrical signals. Found "
            "one good new signal: a failing truck's voltage recovers slowly after engine start. "
            "Raised warning coverage from 5 to 6 of the 10 failures with zero false alarms on "
            "healthy trucks, and made one warning 30 days earlier.")
    _p(doc, "V11.1 — asked the obvious follow-up: if we feed those new warning signs into the "
            "failure-date math (like a doctor adding blood-test results to age), can we finally "
            "predict each truck's date? The answer is NO — and now it is proven, not assumed. "
            "Adding the signals made predictions slightly WORSE (149–162 vs 140 days error). One "
            "trap explained much of it: counting a truck's lifetime 'slow recovery' events mostly "
            "measures how long it has lived, not how sick it is. The question is now closed with "
            "strict no-peeking-into-the-future checks (all five honesty gates passed).")

    _h(doc, "3. What you actually get (deployable today)")
    _bullets(doc, [
        "Risk ranking of in-service trucks (frozen classifier, AUROC 0.927).",
        "Fleet replacement window: plan around ~601 days / ~120,000 km.",
        "RED alarm — electrical fault-code storm (GED=2): fired for 2 of 10 failures "
        "(21-day and 1-day leads), never on a healthy truck.",
        "YELLOW alarm (new in V11.1) — truck CURRENTLY showing the combined electrical-distress "
        "signature: 3 of 10 failures, never on a healthy truck. Feeds the maintenance "
        "recommendation directly.",
        "A guardrail finding: always-on versions of the new alarms would cry wolf (up to 6 of 15 "
        "healthy trucks falsely flagged) — only the 'distressed right now' version is deployed.",
    ])

    # ----- metrics -----
    _h(doc, "4. Metrics — three versions side by side")
    _p(doc, "Shared foundations (identical by design):", bold=True)
    _table(doc,
           ["Metric", "V10.6.2", "V11", "V11.1"],
           [
               ["Classifier AUROC (WHICH truck)", "0.9267 (frozen)", "0.9267 (frozen)", "0.9267 (frozen)"],
               ["Fleet wear-out window", "601 d empirical", "601 d (by reference)", "601 d (by reference)"],
               ["Weibull fit (shape / scale)", "5.17 / 771 d", "— (not re-run)", "5.17 / 771 d (reproduced)"],
           ])

    _p(doc, "")
    _p(doc, "Per-truck RUL accuracy (time-rewound LOVO, lower is better):", bold=True)
    _table(doc,
           ["Model", "Day-MAE", "PI coverage", "Verdict"],
           [
               ["Fleet-clock dummy (601-day rule)", "49.7 d", "—", "still unbeaten"],
               ["V10.6.2 model (no covariates)", "141.8 d", "0.867", "loses to dummy"],
               ["V11.1 M0 (no covariates)", "140.4 d", "0.867", "ships (= V10.6.2)"],
               ["V11.1 M1 (+ crank-recovery count)", "148.8 d", "0.833", "worse — exposure-confounded"],
               ["V11.1 M2 (+ compound-vote state)", "162.2 d", "0.800", "worst"],
           ])
    _p(doc, "Verdict: NO_IMPROVEMENT — covariates do not individualize RUL (third structural "
            "confirmation; gate-enforced honest fallback).", italic=True, size=9.5)

    _p(doc, "")
    _p(doc, "Warning / lead-time channels (failed-truck recall, healthy-truck false alarms):", bold=True)
    _table(doc,
           ["Channel", "V10.6.2", "V11", "V11.1"],
           [
               ["Strict-gate precursor recall", "5/10", "6/10 (0/15 FA)", "carried from V11"],
               ["GED=2 storm (RED alarm)", "2/10, 0/15 FA", "2/10, 0/15 FA", "2/10, 0/15 FA"],
               ["Compound early-watch (YELLOW)", "—", "4/10, 0/15 FA (window test)", "3/10, 0/15 FA (current-state, deployable)"],
               ["Decision matrix", "2-D (risk × time)", "—", "3-D (risk × time × emergency)"],
               ["Best new signal", "—", "post-crank voltage recovery", "same, as alarm input"],
           ])

    _h(doc, "5. Bottom line")
    _bullets(doc, [
        "Same as before: which trucks are risky, and the ~601-day fleet replacement window.",
        "Newly proven: individual failure dates are impossible with this data — stop chasing them.",
        "Newly gained: a second zero-false-alarm warning light and a maintenance decision table "
        "that uses it.",
    ])

    _h(doc, "Appendix — number sources", level=2)
    _table(doc,
           ["Figure", "Source file"],
           [
               ["AUROC 0.9267", "V5.2_ALT/results/V10.5.3_20_5_ALT_final_report.json"],
               ["601 d window", "V10.6.2_ALT/cache/rul/fleet_window.json"],
               ["MAE 141.8 / dummy 49.7", "V10.6.2_ALT/cache/backtest/backtest_results.json"],
               ["MAE 140.4 / 148.8 / 162.2", "V11.1_ALT/cache/backtest/backtest_results.json"],
               ["Recall 5/10 vs 6/10, 0/15 FA", "V11_ALT_heuristics results + cache/forensics"],
               ["Early-watch 3/10, 0/15 FA", "V11.1_ALT/cache/emergency/emergency_per_vin.csv"],
               ["Five gates PASS", "V11.1_ALT/results/V11.1_ALT_verification.json"],
           ], widths=[2.6, 4.2])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"[summary docx] wrote {OUT}")


if __name__ == "__main__":
    main()
