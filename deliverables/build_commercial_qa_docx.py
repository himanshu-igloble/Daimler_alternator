# -*- coding: utf-8 -*-
"""
Build the DICV / BharatBenz "Truck Connect Predictive Intelligence" commercial
Q&A deliverable as ONE polished, fully-editable .docx.

Deployment reality (this revision):
  * Models are TRAINED on ~500 vehicles per vertical  -> transformer tier (A100).
  * Models are DEPLOYED on ~100,000 vehicles in DICV Truck Connect.
  * Production compute is given per 25k vehicle SLAB (25 / 50 / 75 / 100k) and
    per vehicle, with a per-vertical split.
  * Retraining is QUARTERLY only (complete individual models, champion/challenger).
  * Operating rhythm (daily / weekly / monthly / quarterly) is shown explicitly.

House style + helpers reused from build_scaleup_docx.py (navy headers, green rows).
FX: Rs 94 / USD.  Pricing snapshot: 24-Jun-2026, Azure Central India (PAYG / Spot).
All figures are planning estimates for decision-making, NOT a quotation.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import build_scaleup_docx as base  # noqa: E402  (helpers; no file write on import)

from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
FNAME = "DICV_TruckConnect_Predictive_Intelligence_Commercial_QA.docx"
DATE = "2026-06-26"
FX = 94.0  # Rs / USD (single source of truth)

# palette aliases
NAVY, NAVY2 = base.NAVY, base.NAVY2
BLUE, GREEN, AMBER = base.BLUE, base.GREEN, base.AMBER
LGREY, MGREY, WHITE = base.LGREY, base.MGREY, base.WHITE
CENTER, LEFT = base.CENTER, base.LEFT
GREENROW = "C6E0B4"
CALLOUT = "FBE4D5"
INTERNAL = "FCE4EC"
GREY = "595959"


# ----------------------------------------------------------------- money helpers
def rs(usd):
    return f"Rs {round(usd * FX):,}"


def dr(usd):
    return f"${usd:g}  ({rs(usd)})"


# ----------------------------------------------------------------- local helpers
def tline(doc, text, size, color=None, bold=True, align=CENTER, space=2, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size); r.font.name = 'Calibri'
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def q(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Q{num}.  {text}")
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor.from_string(NAVY)
    return p


def a(doc, text, space=6, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r0 = p.add_run("A.  ")
    r0.bold = True; r0.font.size = Pt(size); r0.font.name = 'Calibri'
    r0.font.color.rgb = RGBColor.from_string(GREEN)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = 'Calibri'
    return p


def cont(doc, text, space=6, size=10.5, italic=False, bold=False):
    return base.para(doc, text, size=size, italic=italic, space=space, bold=bold)


def abullet(doc, label, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    if label:
        rl = p.add_run(label + " ")
        rl.bold = True; rl.font.size = Pt(size); rl.font.name = 'Calibri'
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = 'Calibri'
    return p


def callout(doc, text, fill=CALLOUT, color="833C00", size=9.5, bold=False, width=6.9):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    base.fixed_layout(t)
    c = t.rows[0].cells[0]
    c.width = Inches(width)
    base.shade(c, fill); base.vcenter(c)
    base.set_cell(c, text, bold=bold, size=size, color=color)
    return t


def spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p


def page_field(p):
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'PAGE'
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(b); r.append(i); r.append(s); r.append(e)


def setup_footer(doc):
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.alignment = CENTER
    r = p.add_run("ByteEdge  -  DICV Truck Connect Predictive Intelligence  -  planning estimates, not a quotation  -  page ")
    r.font.size = Pt(8); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor.from_string(GREY)
    page_field(p)


# ============================================================================ build
def build():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
        sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10.5)

    # --------------------------------------------------------------- title
    tline(doc, "DICV  |  BharatBenz Predictive Maintenance", 11, color=GREY, bold=True, space=2)
    tline(doc, "Daimler Truck Connect - Predictive Intelligence", 20, color=NAVY, bold=True, space=2)
    tline(doc, "Commercial Q&A  &  Costing Annex", 15, color=NAVY2, bold=True, space=4)
    tline(doc, "Trained on ~500 vehicles / vertical (transformer)  -  deployed on ~100,000 vehicles in Truck Connect",
          10.5, color=GREY, bold=False, italic=True, space=2)
    tline(doc, f"Prepared by ByteEdge   -   {DATE}   -   FX Rs 94/USD   -   pricing snapshot 24-Jun-2026 (Azure Central India)",
          9, color=GREY, bold=False, space=8)

    callout(doc,
            "Customer-facing answers are in Sections 1-8. Annex A shows the compute as (hours x per-hour rate); Annex B "
            "holds internal commercial notes (cost-to-price bridge, margin %) and is NOT for external distribution. "
            "Planning estimates for decision-making, not a quotation.",
            fill=LGREY, color="1F3864", size=9.5)
    spacer(doc, 6)

    # --------------------------------------------------------------- AT A GLANCE
    base.h(doc, "At a glance", 1)

    base.para(doc, "Production compute by deployment size (daily batch scoring, all 3 verticals):", bold=True, size=10.5)
    base.add_table(doc,
        ["Fleet (once-daily batch)", "Compute / month", "Compute / year", "Per vehicle / year"],
        [["25,000 vehicles", "~Rs 0.55-0.85 L", "~Rs 6.5-10.3 L", "~Rs 26-41"],
         ["50,000 vehicles", "~Rs 0.8-1.3 L", "~Rs 9.8-15.3 L", "~Rs 20-31"],
         ["75,000 vehicles", "~Rs 1.0-1.6 L", "~Rs 12.4-19.3 L", "~Rs 16-26"],
         ["100,000 vehicles (target)", "~Rs 1.2-1.9 L", "~Rs 14.6-22.8 L", "~Rs 15-23"]],
        widths=[2.1, 1.6, 1.7, 1.5], font=9.5, special={3: GREENROW})
    base.para(doc, "Twice-daily scoring ~ 1.5-1.8x the monthly figure. Per-vehicle cost falls as the fleet grows "
                   "(shared platform spread over more trucks).", size=9, italic=True)

    base.para(doc, "Model tier - changes TRAINING only, not the run-rate above:", bold=True, size=10.5)
    base.add_table(doc,
        ["Metric", "Small models (small data)", "Transformer (~500 veh/vertical)"],
        [["Model family", "Ridge / XGBoost / IForest + small LSTM/TCN", "Temporal Fusion Transformer"],
         ["Training GPU", "A10 (24 GB)", "A100 (80 GB)"],
         ["Build / industrialization compute", "~Rs 0.7-1.2 L", "~Rs 1.3-5.0 L (~ Rs 4-4.5 L envelope)"],
         ["Production run-rate & per-vehicle", "as table above", "identical (inference is model-agnostic)"],
         ["Quarterly retrain, 3 verticals", "~Rs 450-1,300 / yr (A10)", "~Rs 5,000-11,500 / yr (A100)"]],
        widths=[2.0, 2.45, 2.45], font=9)
    callout(doc, "Bottom line: at 100,000 vehicles, once-daily scoring costs ~Rs 1.2-1.9 Lakh/month (~Rs 15-23 per "
                 "vehicle/year). The move to a transformer raises only TRAINING / retraining compute (still a few "
                 "thousand rupees/year); the production run-rate is unchanged because daily batch inference of a "
                 "still-small model is cheap (< 15 min/pass on CPU in the consolidated plan).",
            fill=GREENROW, color="1F3864", size=9.5, bold=True)
    doc.add_page_break()

    # --------------------------------------------------------------- TOC
    base.h(doc, "Contents", 1)
    base.add_toc(doc)
    doc.add_page_break()

    # =========================================================== 1
    base.h(doc, "1.  How to read this document", 1)
    abullet(doc, "Customer answers (Sec. 1-8):", "clean, decision-ready.")
    abullet(doc, "Annex A:", "compute as hours x per-hour rate, plus the storage / platform lines that are not hour-priced.")
    abullet(doc, "Annex B (internal):", "cost-to-price bridge, margin %, points to confirm. Remove before sharing externally.")
    base.para(doc, "Currency Rs 94/USD. Cloud prices: Azure Central India, 24-Jun-2026. Deployment ~100,000 vehicles; "
                   "models trained on ~500 vehicles/vertical (transformer tier).", italic=True, size=9.5)

    # =========================================================== 2
    base.h(doc, "2.  Programme & Scope", 1)

    q(doc, 1, "What does the Rs 24 Lakhs “Clutch 500 Readiness” programme cover?")
    a(doc, "It takes the proven clutch model from pilot to a production-grade capability ready for the fleet and "
           "integrated into Daimler Truck Connect, across three workstreams:")
    abullet(doc, "Data Engineering & Foundation -", "production-grade ingestion, cleaning and sentinel-handling, a "
            "curated data store and feature store, plus data-quality and silent-telemetry monitors.")
    abullet(doc, "Predictive Intelligence Industrialization -", "hardening the clutch model (anomaly baseline + "
            "classifier + RUL with conformal intervals), the champion/challenger retraining loop, and validation gates.")
    abullet(doc, "Truck Connect Readiness & Architecture -", "the API contract, webhook alerts, the signed/licensed "
            "container, and the deployment architecture.")
    cont(doc, "The same foundation and Truck Connect plumbing are reused by the alternator and starter verticals - which "
              "is why each industrializes for an incremental Rs 7 Lakhs. Clutch is the most mature (~63-day failure lead).")

    q(doc, 2, "What is the breakup of the Rs 24 Lakhs?")
    base.add_table(doc,
        ["Workstream", "Amount", "What it includes"],
        [["Data Engineering & Foundation", "Rs 7 L", "Ingestion, cleaning, sentinel rules, curated + feature store, data-quality monitors"],
         ["Predictive Intelligence Industrialization", "Rs 13 L", "Clutch model hardening, calibration, champion/challenger loop, LOVO validation, registry"],
         ["Truck Connect Readiness & Architecture", "Rs 4 L", "API contract, webhook alerts, signed licensed container, deployment design"],
         ["Total", "Rs 24 L", "~30% data / 54% predictive-model industrialization / 16% Truck Connect readiness"]],
        widths=[2.5, 0.8, 3.6], font=9, special={3: GREENROW})

    q(doc, 3, "How much compute cost is assumed (build phase)?")
    a(doc, "~Rs 4-4.5 Lakhs across development, training, storage and validation. For small models the actual build "
           "compute is only ~Rs 0.7-1.2 L (the envelope is conservative); for the transformer tier (~500 veh/vertical, "
           "A100, ~600 GPU-hours incl. hyper-parameter search) it is ~Rs 1.3-5.0 L - i.e. the Rs 4-4.5 L envelope still "
           "covers the typical build. Most of the Rs 24 L is engineering effort, not compute.")

    q(doc, 4, "Does the Rs 24 Lakhs cover all tractor-trailer variants?")
    a(doc, "Yes - assuming a similar clutch architecture. Variants sharing the same clutch hardware and signal set are "
           "covered by the same model with minor re-calibration. Materially different hardware, torque ratings or duty "
           "cycles may need incremental calibration - a small add-on, not a fresh programme. A variant-to-clutch-hardware "
           "mapping from Daimler confirms which variants are in scope at no extra cost (Section 8).")

    # =========================================================== 3  DEPLOYMENT COMPUTE
    base.h(doc, "3.  Deployment Compute (100,000 vehicles, by 25k slab)", 1)
    base.para(doc, "Batch-first architecture: a scheduled CPU batch scores the fleet daily; the GPU is summoned only for "
                   "quarterly retraining. Compute is model-agnostic (small models or transformer score equally cheaply "
                   "in batch). Hours x rate in Annex A.", italic=True, size=9.5)

    q(doc, 5, "Production compute by deployment slab (once daily)?")
    base.add_table(doc,
        ["Fleet (once-daily batch)", "Compute / month", "Compute / year", "Per vehicle / year"],
        [["25,000 vehicles", "~Rs 0.55-0.85 L", "~Rs 6.5-10.3 L", "~Rs 26-41"],
         ["50,000 vehicles", "~Rs 0.8-1.3 L", "~Rs 9.8-15.3 L", "~Rs 20-31"],
         ["75,000 vehicles", "~Rs 1.0-1.6 L", "~Rs 12.4-19.3 L", "~Rs 16-26"],
         ["100,000 vehicles (target)", "~Rs 1.2-1.9 L", "~Rs 14.6-22.8 L", "~Rs 15-23"]],
        widths=[2.1, 1.6, 1.7, 1.5], font=9.5, special={3: GREENROW})
    cont(doc, "Compute scales sub-linearly with fleet size (storage and per-vehicle CPU grow ~linearly; the platform / "
              "monitoring layer is largely fixed), so per-vehicle cost falls from ~Rs 26-41 at 25k to ~Rs 15-23 at 100k.",
         size=10)

    q(doc, 6, "Production compute if scoring runs twice daily?")
    a(doc, "~1.5-1.8x the once-daily figure (e.g. ~Rs 1.8-3.4 Lakh/month at 100k). It does not double: telemetry is "
           "ingested and stored once; only the second scoring pass, the extra prediction-history writes and pipeline runs "
           "add cost. The operating rhythm (Section 4) is once-daily, so once-daily is the planning figure.")

    q(doc, 7, "Per-vehicle compute cost, split by vertical (at 100k)?")
    a(doc, "All three verticals together cost ~Rs 15-23/vehicle/year. The split is roughly even because most of the cost "
           "is shared infrastructure (storage, feature pipeline, platform); the model inference itself is negligible:")
    base.add_table(doc,
        ["Vertical", "Daily output", "Compute (Rs / veh / yr @ 100k)"],
        [["Clutch", "RUL + risk zone (anomaly-based, ~63-day lead)", "~Rs 7-9"],
         ["Alternator", "Risk zone + maintenance window + GED=2 alarm", "~Rs 5-7"],
         ["Starter", "Risk zone + window + persistent-RED dwell pager", "~Rs 5-7"],
         ["All three (shared)", "combined daily card per vehicle", "~Rs 15-23"]],
        widths=[1.5, 3.8, 1.6], font=9, special={3: GREENROW})
    cont(doc, "Clutch is marginally higher (anomaly + RUL). Adding a vertical is incremental (~Rs 5-7/veh/yr), not a "
              "third of the cost. This is compute cost, distinct from the SaaS price in Section 7.", size=9.5, italic=True)

    # =========================================================== 4  OPERATING RHYTHM
    base.h(doc, "4.  Operating Rhythm", 1)
    base.para(doc, "One cadence runs the deployed system; retraining sits at the bottom of it, gated by new labels.")
    base.add_table(doc,
        ["Cadence", "Activity", "Output", "Honest note"],
        [["Daily", "Batch inference", "Risk tier + window + alarms (RUL for clutch)", "per-component split below"],
         ["Weekly", "Aggregate daily scores", "Predictive-maintenance calendar / depot queue", "dwell pager runs on weekly windows"],
         ["Monthly", "Performance review", "Drift, FP/FN vs service outcomes, stale-truck audit", "flag silent telemetry, not false-GREEN"],
         ["Quarterly", "Retrain if enough new labels", "Champion / challenger promotion", "trigger = new failures, not the date"]],
        widths=[0.95, 1.5, 2.55, 1.9], font=9)
    base.para(doc, "Daily output by component:", bold=True, size=10)
    base.add_table(doc,
        ["Component", "Daily inference output", "Is it an RUL?"],
        [["Clutch", "RUL estimate + risk zone (anomaly-based, ~63-day lead)", "Yes - RUL (caveated)"],
         ["Alternator", "Risk zone + maintenance window + GED=2 disturbance alarm", "No - risk zone"],
         ["Starter", "Risk zone + maintenance window + persistent-RED dwell pager (weekly)", "No - risk zone"]],
        widths=[1.2, 4.3, 1.4], font=9)

    # =========================================================== 5  RETRAINING
    base.h(doc, "5.  Retraining (quarterly)", 1)
    q(doc, 8, "Retraining compute - complete individual models, quarterly, on Azure (all 3 verticals)?")
    a(doc, "Retraining is quarterly only (champion/challenger, triggered by new confirmed failures - not the calendar). "
           "It is periodic retraining, not reinforcement learning. Cost depends on the model tier and is fleet-size-"
           "independent (the training matrix is small regardless of deployment size):")
    base.add_table(doc,
        ["Model tier", "GPU", "Per quarterly run (3 verticals)", "Per year (4 quarters)"],
        [["Small models", "A10 Spot", f"1-3 GPU-h  (~{rs(1.0)}-{rs(3.5)})", "~Rs 450-1,300"],
         ["Transformer (~500 veh/vertical)", "A100 Spot", f"12-30 GPU-h  (~{rs(13.4)}-{rs(30.5)})", "~Rs 5,000-11,500"]],
        widths=[2.4, 1.0, 2.0, 1.5], font=9, special={1: GREENROW})
    cont(doc, "Per vertical, the transformer quarterly retrain is ~Rs 1,700-3,800/year. Even on dedicated (PAYG) GPUs the "
              "all-vertical figure stays under ~Rs 60k/year. The real constraint is how many new labelled failures arrive, "
              "not the compute. Hours x rate in Annex A.3.", size=10)

    # =========================================================== 6  EXPANSION
    base.h(doc, "6.  Expansion Strategy", 1)
    q(doc, 9, "Recommended expansion strategy?")
    a(doc, "Industrialize by component first - Clutch -> Alternator -> Starter -> Brake -> Battery - then extend across "
           "vehicle families. The data foundation and Truck Connect integration are built once and amortized, so each "
           "additional component is an incremental Rs 7 Lakhs; extending to other families is mainly re-calibration.")
    q(doc, 10, "Single-Unit Truck clutch expansion?")
    a(doc, "~10-12 weeks, Rs 10-12 Lakhs. Re-uses the proven clutch model and existing foundation; re-maps and "
           "re-calibrates features for the single-unit duty cycle, re-validates, and onboards into Truck Connect.")
    q(doc, 11, "Mining / Off-Highway clutch expansion?")
    a(doc, "~12-16 weeks, Rs 12-15 Lakhs. Off-highway duty cycles differ substantially (severe loads, low-speed/"
           "high-torque, harsh thermal/dust), needing more re-calibration, possible new failure modes and longer validation.")
    q(doc, 12, "Alternator industrialization cost?")
    a(doc, "Rs 7 Lakhs incremental on the shared foundation: validated alternator model (AUROC 0.927) to production - "
           "GED-disturbance alarm, risk-window matrix, calibration, champion/challenger. Deliverable = risk score + "
           "maintenance window + GED alarm.")
    q(doc, 13, "Starter industrialization cost?")
    a(doc, "Rs 7 Lakhs incremental: validated starter model (AUROC ~0.93) to production with its maintenance window and a "
           "persistent-RED dwell pager. Deliverable = risk zone + window + alarm, not a per-day countdown.")

    # =========================================================== 7  PRICING
    base.h(doc, "7.  Commercial Model & Pricing", 1)
    q(doc, 14, "Recommended commercial model?")
    a(doc, "Offer both an Enterprise License (large committed / OEM-embedded fleets; capacity-based) and a Per-Vehicle "
           "SaaS subscription (variable/growing fleets; transparent Rs/vehicle/year). Both lets Daimler choose by "
           "deployment model and captures large-fleet commitments plus incremental adoption.")
    q(doc, 15, "Recommended enterprise pricing?")
    base.add_table(doc,
        ["Fleet size band", "Enterprise price (full suite)"],
        [["Up to 25,000 vehicles", "Rs 1.5 Crore"],
         ["25,000 - 50,000 vehicles", "Rs 2.5 Crore"],
         ["50,000 - 100,000 vehicles", "Rs 3.5 Crore"],
         ["Above 100,000 vehicles", "Custom"]],
        widths=[3.4, 3.5], font=10)
    q(doc, 16, "Recommended per-vehicle pricing?")
    base.add_table(doc,
        ["Component bundle", "Price (per vehicle / year)"],
        [["Clutch", "Rs 150"],
         ["Clutch + Starter", "Rs 250"],
         ["Clutch + Starter + Alternator", "Rs 400"],
         ["Full Suite", "Rs 600"]],
        widths=[3.9, 3.0], font=10)
    cont(doc, "Pricing reflects avoided breakdowns and downtime; the underlying cloud compute is only ~Rs 15-23/vehicle/"
              "year (Section 3) - a small fraction of price.", size=9.5, italic=True)

    # =========================================================== 8  NEEDS
    base.h(doc, "8.  What we need from Daimler", 1)
    q(doc, 17, "What information is still needed?")
    abullet(doc, "Vehicle variant mapping -", "families/variants in scope and their clutch (and alternator/starter) hardware.")
    abullet(doc, "Deployment architecture -", "Truck Connect integration contract (APIs, auth, latency/SLA); cloud / on-prem / hybrid.")
    abullet(doc, "Cloud preference -", "Azure (assumed) and target region / data-residency (Central India assumed).")
    abullet(doc, "IT / security -", "data residency, RBAC, audit logging, per-tenant isolation, signed-image / license controls.")
    abullet(doc, "Prediction frequency & retention -", "once- vs twice-daily scoring (sets compute per Section 3) and "
            "prediction-history retention.")

    # =========================================================== ANNEX A
    doc.add_page_break()
    base.h(doc, "Annex A.  Compute - hours x rate", 1)
    base.para(doc, "Only GPU and CPU are hour-priced. Storage is per GB-month and platform/monitoring are flat fees - "
                   "together about half the bill - so those are shown in their real units.", italic=True, size=9.5)

    base.h(doc, "A.1  Rate card (Central India, 24-Jun-2026)", 2)
    base.add_table(doc,
        ["Resource", "Unit", "PAYG rate", "Spot rate"],
        [["A10 GPU (NV36ads_A10_v5)", "$/GPU-hr", dr(4.48), dr(0.828)],
         ["A100 GPU (NC24ads_A100_v4)", "$/GPU-hr", dr(5.142), dr(0.950)],
         ["D16ds_v5 CPU (16 vCPU - feature/ETL)", "$/instance-hr", dr(0.918), "~" + dr(0.30)],
         ["D8ds_v5 CPU (8 vCPU - inference)", "$/instance-hr", dr(0.459), "~" + dr(0.15)],
         ["Blob storage hot / cool / archive", "$/GB-month", "$0.020 / 0.010 / 0.002", "-"]],
        widths=[2.6, 1.3, 1.6, 1.4], font=9)

    base.h(doc, "A.2  Once-daily production @ 100,000 vehicles", 2)
    base.para(doc, "Hour-based (compute), per month:", bold=True, size=10)
    base.add_table(doc,
        ["Job", "Instance", "Instance-hr/mo", "= $/mo PAYG", "Spot option"],
        [["Feature aggregation", "D16ds_v5", "435-1,307", "$400-1,200", "$130-390"],
         ["Batch inference (x3 verticals)", "D8ds_v5", "218-654", "$100-300", "$33-98"],
         ["GPU retrain (quarterly, amortized)", "A100 Spot", "~4-10 GPU-hr", "-", "$4-10"],
         ["Compute subtotal", "", "~650-1,960 CPU-hr/mo", "$500-1,500", "$165-500"]],
        widths=[2.0, 0.9, 1.4, 1.4, 1.2], font=8.5, special={3: MGREY})
    base.para(doc, "Not hour-based:", bold=True, size=10)
    base.add_table(doc,
        ["Line", "Quantity", "= $/mo"],
        [["Storage (raw 14-day hot + curated + features)", "~12,000 GB-mo @ $0.02/GB-mo", "$220-500"],
         ["Registry + orchestration + monitoring", "platform (flat) + per-run", "$200-520"]],
        widths=[3.6, 2.3, 1.0], font=9)
    callout(doc, "Total once-daily @ 100k: ~$950-2,560/mo = ~Rs 0.9-2.4 lakh; planning figure ~Rs 1.2-1.9 Lakh/month. "
                 "Compute is model-agnostic - a transformer scores just as cheaply in daily batch (< 15 min/pass on CPU).",
            fill=GREENROW, color="1F3864", size=9.5, bold=True)

    base.h(doc, "A.3  Retraining (quarterly) - hours x rate", 2)
    base.para(doc, "One quarterly retrain of all three verticals; transformer tier on A100 (small-model tier on A10):", size=10)
    base.add_table(doc,
        ["Tier / GPU", "Hours", "= Spot / run", "= PAYG / run", "Per year (x4) Spot"],
        [["Small / A10", "1-3 GPU-h", f"{rs(1.0)}-{rs(3.5)}", f"{rs(4.7)}-{rs(14.4)}", "Rs 450-1,300"],
         ["Transformer / A100", "12-30 GPU-h", f"{rs(13.4)}-{rs(30.5)}", f"{rs(63.7)}-{rs(156.3)}", "Rs 5,000-11,500"]],
        widths=[1.9, 1.1, 1.3, 1.3, 1.3], font=9, special={1: GREENROW})
    abullet(doc, "Quarterly only:", "4 runs/year. GPU-hours/year = 48-120 (transformer) or 4-12 (small). Per vertical, "
            "the transformer is ~Rs 1,700-3,800/year.")
    abullet(doc, "Fleet-size-independent & checkpointed:", "same cost for 25k or 100k vehicles; Spot eviction just "
            "resumes, so PAYG is only the rare must-finish ceiling.")

    base.h(doc, "A.4  Bottom line", 2)
    base.para(doc, "Compute hours are nearly free: ~1-3 CPU-seconds per vehicle per day to score, and 48-120 GPU-hours "
                   "per year to retrain all three transformers. About half the production bill is storage (GB-month) and "
                   "flat platform fees, so the real cost levers are architecture (batch vs ~4x costlier always-on "
                   "streaming) and retention - not the hours.", bold=True, size=10)

    # =========================================================== ANNEX B
    doc.add_page_break()
    base.h(doc, "Annex B.  Internal commercial notes", 1)
    callout(doc, "INTERNAL - NOT FOR EXTERNAL DISTRIBUTION. Remove before sharing with the customer.",
            fill=INTERNAL, color="880E4F", size=9.5, bold=True)
    spacer(doc, 4)

    base.h(doc, "B.1  Cost-to-price bridge", 2)
    base.para(doc, "Consolidated 500-truck build (all components + data + Truck Connect + onboarding) ~ $60.7-66.7k = "
                   "~Rs 57-63 Lakhs at Rs 94/USD. Rs 24 L clutch + Rs 7 L alternator + Rs 7 L starter = Rs 38 L covers "
                   "the foundation + 3 verticals + Truck Connect readiness; the remaining ~Rs 19-25 L is brakes "
                   "(greenfield), onboarding and the fuller Truck Connect integration. The offer is internally consistent.")

    base.h(doc, "B.2  Effort reconciliation ($40/hr fully-loaded, Rs 94/USD)", 2)
    base.add_table(doc,
        ["Item", "Price", "~ eng-hours", "~ eng-days"],
        [["Rs 24 L clutch programme", "Rs 24 L", "~640", "~80"],
         ["Alternator / Starter (each)", "Rs 7 L", "~185", "~23"],
         ["Single-Unit Truck clutch", "Rs 10-12 L", "~265-320", "~33-40"],
         ["Mining / Off-Highway clutch", "Rs 12-15 L", "~320-400", "~40-50"]],
        widths=[3.0, 1.3, 1.3, 1.3], font=9, special={0: MGREY})

    base.h(doc, "B.3  Points to confirm before quoting", 2)
    abullet(doc, "1. Cost vs billing rate.", "$40/hr is fully-loaded (cost). Rs 24 L ~ 640 hours at that rate = near "
            "delivery cost; confirm whether $40/hr is cost or billing and set the margin before quoting.")
    abullet(doc, "2. Enterprise vs per-vehicle.", "Rs 1.5 Cr / 25,000 = Rs 600/veh = the full-suite per-vehicle rate. "
            "Decide annual vs perpetual; if perpetual, add an AMC (~18-22%/yr) for compute, support and retraining.")
    abullet(doc, "3. Starter vs alternator effort.", "Starter is a 6-week plan phase vs alternator 4; both at Rs 7 L is "
            "fine if pilot reuse offsets it, else consider Rs 8-9 L for starter.")
    abullet(doc, "4. Architecture dependency.", "The Rs 1.2-1.9 L/mo @ 100k assumes the batch architecture. Real-time / "
            "always-on serving moves it toward the heavier streaming profile - hence the prediction-frequency question.")
    abullet(doc, "5. Deliverable scope.", "Clutch = RUL (~63-day lead). Alternator & starter = risk zone + window + "
            "alarm, NOT a per-day RUL. Do not sell daily RUL for alternator/starter.")
    abullet(doc, "6. Model tier is priced in.", "The effort prices assume DL/transformer work; the transformer tier "
            "raises only compute (build + retrain), both within the quoted envelopes - not the engineering price.")

    base.h(doc, "B.4  Compute as % of price (margin sanity)", 2)
    base.add_table(doc,
        ["Enterprise band", "License", "Compute / yr (at top of band)", "Compute %"],
        [["Up to 25,000", "Rs 1.5 Cr", "Rs 6.5-10.3 L", "4.3-6.9%"],
         ["25,000-50,000", "Rs 2.5 Cr", "Rs 9.8-15.3 L", "3.9-6.1%"],
         ["50,000-100,000", "Rs 3.5 Cr", "Rs 14.6-22.8 L", "4.2-6.5%"]],
        widths=[1.9, 1.5, 2.0, 1.5], font=9)
    base.para(doc, "Per-vehicle: compute is ~Rs 15-41/veh/yr (fleet-dependent) against Rs 150-600 price -> ~3-7%. Healthy "
                   "infra margin in both models; transformer retraining adds only ~Rs 5-12k/year fleet-wide, immaterial "
                   "to the ratios.", italic=True, size=9.5)

    # =========================================================== assumptions
    doc.add_page_break()
    base.h(doc, "Assumptions, sources & caveats", 1)
    abullet(doc, "FX:", "Rs 94 / USD.")
    abullet(doc, "Pricing snapshot:", "24-Jun-2026, Azure Central India, Linux PAYG / Spot - re-check the Azure calculator before commitment.")
    abullet(doc, "Architecture:", "batch-first (daily CPU batch scoring, on-demand Spot GPU for quarterly retraining); not always-on streaming.")
    abullet(doc, "Model tiers:", "small models (A10) for small data; transformer (A100) at ~500 veh/vertical. Production inference is daily CPU batch in both - run-rate is model-agnostic.")
    abullet(doc, "Deployment / training scale:", "deployed ~100,000 vehicles; trained ~500 vehicles/vertical.")
    abullet(doc, "Scope of compute figures:", "cloud compute + data platform only. Excludes one-time build, dashboard hosting, support, Truck Connect integration engineering and data egress.")
    abullet(doc, "Sources:", "DICV 500-Truck Scale-Up Plan (Consolidated) and the Azure Deployment / Retraining / GPU plans (v2 batch deployment supersedes v1).")
    abullet(doc, "Status:", "planning estimates, NOT a quotation.")
    spacer(doc, 6)
    tline(doc, "DICV BharatBenz Predictive Maintenance  -  Truck Connect Predictive Intelligence  -  Commercial Q&A  -  " + DATE,
          9, color=GREY, bold=False, space=2)

    setup_footer(doc)
    out = os.path.join(OUT_DIR, FNAME)
    doc.save(out)
    print("WROTE:", out)
    print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))


if __name__ == "__main__":
    build()
