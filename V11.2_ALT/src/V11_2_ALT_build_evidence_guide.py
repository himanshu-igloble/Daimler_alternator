"""
V11.2_ALT — build the per-VIN Evidence-Stack Reading Guide (Markdown + DOCX).

For each of the 25 VINs it extracts the truck's actual facts from build_bundle
(M5 zone transitions + drivers, precursor breaches + lead, GED2 storm, charging-
ceiling delta, clock-vs-physics), classifies an archetype, and authors a SIMPLE +
TECHNICAL reading plus KEY INSIGHTS — all grounded in the numbers. Embeds each
evidence-stack figure. Run with: py -3
"""
from __future__ import annotations
import os, sys, json
import numpy as np
import pandas as pd

ROOT = r"D:/Daimler-starter_motor_alternator_battery"
sys.path.insert(0, os.path.join(ROOT, "V11.2_ALT", "src"))
import V11_2_ALT_rul_evidence_stack as E
import V11_2_ALT_rul_evidence_stack_snapshot as S

VIZ_REL = "../visualizations/rul_evidence_stack"
VIZ_ABS = os.path.join(ROOT, "V11.2_ALT", "visualizations", "rul_evidence_stack")
REPORTS = os.path.join(ROOT, "V11.2_ALT", "reports")
RESULTS = os.path.join(ROOT, "V11.2_ALT", "results")
DATE = "2026-06-24"

ALL = [f"VIN{i}_F_ALT" for i in range(1, 11)] + [f"VIN{i}_NF_ALT" for i in range(1, 16)]


# ---------------------------------------------------------------- fact extraction
def clock_zone(rul):
    return "GREEN" if rul > 180 else "YELLOW" if rul > 90 else "ORANGE" if rul > 30 else "BLACK"


def ceiling(bins):
    c, m = bins
    if len(m) == 0:
        return float("nan")
    sel = c >= 1500
    return float(np.median(m[sel])) if sel.any() else float(np.median(m))


def facts(vin):
    b = E.build_bundle(vin)
    breaches = []
    for col, info in b["prec"].items():
        if info["breach"] is not None:
            lead = (b["end"] - pd.Timestamp(b["daily"]["date"].values[info["breach"]])).days
            breaches.append((info["label"], int(lead)))
    breaches.sort(key=lambda t: -t[1])
    ged = bool(b["ged_fired"])
    ged_lead = b["ged_lead"] if ged else None
    early_c, late_c = ceiling(b["chg"]["early"]), ceiling(b["chg"]["late"])
    cdelta = (late_c - early_c) if (np.isfinite(early_c) and np.isfinite(late_c)) else float("nan")
    m5_final = E.zone_of_m5(float(b["m5"][-1]))
    f = dict(
        vin=b["vin"], dname=b["dname"], failed=b["failed"], risk_band=b["risk_band"],
        est_km=b["est_km"], median_rul=b["median_rul"],
        fw_p25=b["fw_p25"], fw_p75=b["fw_p75"], fleet_med=b["fleet_med"],
        rul_last=float(b["rul_med"][-1]), clock_zone_last=clock_zone(float(b["rul_med"][-1])),
        gap=((b["jc"] - b["t1"]).days if b["failed"] else None),
        over_pred=((float(b["rul_med"][-1]) - (b["jc"] - b["t1"]).days) if b["failed"] else None),
        m5_trans=[(t["date"].strftime("%b %Y"), t["to"], t["driver"]) for t in b["m5_trans"]],
        m5_final=m5_final, m5_peak=b.get("m5_peak_transient"),
        breaches=breaches, ged=ged, ged_lead=ged_lead,
        cdelta=cdelta, n_cross=len(b["rul_cross"]),
        cross=[(c["name"], c["corrob"], (c["breaching"][0][0] if c["breaching"] else ("GED2 storm" if c["ged"] else None))) for c in b["rul_cross"]],
    )
    # archetype (precursors = signals genuinely elevated at failure / now)
    if f["failed"]:
        if ged:
            f["arch"] = f"Early-warning (GED2 storm ~{ged_lead}d)"
        elif breaches:
            f["arch"] = f"Late precursor: {breaches[0][0]} (~{breaches[0][1]}d)"
        else:
            f["arch"] = "Silent / abrupt failure"
    else:
        f["arch"] = "In-service — watch" if breaches else "In-service — quiet"
    f["snap"] = S.compute_snapshot(b)   # component × key-moment snapshot table
    return f


# ---------------------------------------------------------------- prose authoring
def author_simple(f):
    n = f["dname"]
    if f["failed"]:
        warn = ("It gave an early electrical warning: the {0} signal stood out about {1} days before it failed."
                .format(f["breaches"][0][0], f["breaches"][0][1]) if f["breaches"]
                else "It gave no early electrical warning — it failed quietly, with the sensors looking normal until the end.")
        ged = (f" The excitation-fault 'storm' also fired about {f['ged_lead']} days ahead — a rare, high-confidence alarm." if f["ged"] else "")
        chg = ("Its charging voltage held steady to the end." if not np.isfinite(f["cdelta"]) or f["cdelta"] > -0.3
               else f"Its charging voltage sagged about {abs(f['cdelta']):.1f} V late in life.")
        over = (f" Critically, the fleet schedule still expected ~{f['rul_last']:.0f} more days at that point — it "
                f"**over-predicted by ~{f['over_pred']:.0f} days** because this truck failed early; that gap is why we "
                f"report a fleet window, not a per-truck day-count." if f.get("over_pred") is not None else "")
        return (f"**{n}** is a truck that **failed**. On the fleet's wear-out clock it still looked like it had "
                f"~{f['rul_last']:.0f} days of life at last contact (clock zone **{f['clock_zone_last']}**). {warn}{ged} {chg}{over}")
    else:
        watch = (f"its **{f['breaches'][0][0]}** signal is currently running high versus its own baseline — worth a check"
                 if f["breaches"] else "all its signals are calm")
        return (f"**{n}** is **still in service**. The fleet clock estimates roughly **{f['median_rul']:.0f} days** of "
                f"life remaining. Right now {watch} — routine monitoring against the ~{f['fleet_med']:.0f}-day fleet window.")


def author_technical(f):
    parts = [f"Risk band **{f['risk_band']}**; est. ~{f['est_km']/1000:.0f}k km; fleet wear-out window "
             f"{f['fw_p25']:.0f}–{f['fw_p75']:.0f} d (median {f['fleet_med']:.0f})."]
    if f["m5_trans"]:
        tr = "; ".join(f"{d}→{z} (driver: {drv})" for d, z, drv in f["m5_trans"])
        parts.append(f"Parametric M5 health-zone transitions: {tr}; final M5 zone **{f['m5_final']}**.")
    else:
        parts.append(f"M5 health-zone stayed **{f['m5_final']}** throughout (no parametric transition).")
    if f["breaches"]:
        if f["failed"]:
            parts.append("Signals genuinely elevated at failure (σ≥2 at the end): "
                         + ", ".join(f"{lab} (run-up ~{lead} d)" for lab, lead in f["breaches"]) + ".")
        else:
            parts.append("Currently-elevated signals (σ≥2 at last data): "
                         + ", ".join(f"{lab} (~{lead} d run)" for lab, lead in f["breaches"]) + ".")
    else:
        parts.append("No signal is elevated at the end (σ<2) — no genuine precursor.")
    if f.get("m5_peak"):
        parts.append(f"M5 transiently peaked at {f['m5_peak']} then recovered (not a sustained zone change).")
    if f["failed"] and f.get("over_pred") is not None:
        parts.append(f"Schedule over-prediction: fleet clock RUL@last {f['rul_last']:.0f} d vs {f['gap']} d to actual "
                     f"failure → over by ~{f['over_pred']:.0f} d (clock zone {f['clock_zone_last']} at failure).")
    if f["ged"]:
        parts.append(f"GED2 excitation storm fired (~{f['ged_lead']} d lead).")
    if np.isfinite(f["cdelta"]):
        parts.append(f"Charging ceiling (≥1500 rpm) early→late Δ = {f['cdelta']:+.2f} V "
                     f"({'ceiling collapse' if f['cdelta'] <= -0.3 else 'stable regulation'}).")
    if f["failed"]:
        if f["n_cross"] == 0:
            parts.append(f"RUL clock never crossed a lower zone before failure — it failed while the schedule still read **{f['clock_zone_last']}** (clock ≠ physics).")
        else:
            cs = "; ".join(f"{nm} ({'physics✓ '+sig if cor else 'schedule-only'})" for nm, cor, sig in f["cross"])
            parts.append(f"RUL-clock crossings: {cs}.")
    return " ".join(parts)


def author_insights(f):
    out = []
    if f["failed"]:
        if f["ged"]:
            out.append(f"**Actionable signature.** The GED2 storm gives ~{f['ged_lead']} d notice — the one archetype where a short-horizon pager is viable; deploy the GED channel for trucks like this.")
        if f["breaches"]:
            out.append(f"**Physics led the schedule.** {f['breaches'][0][0]} deviated ~{f['breaches'][0][1]} d before failure — earlier evidence than the fleet clock provided.")
        if not f["breaches"] and not f["ged"]:
            out.append("**Hardest case — silent failure.** No precursor and no parametric zone change; only the fleet replacement window protects this truck. Confirms the abrupt-failure archetype.")
        if f["n_cross"] == 0:
            out.append(f"**Clock ≠ physics.** It failed while the RUL schedule still read {f['clock_zone_last']} — the parameters (M5 / precursors), not the calendar, are what flagged the risk.")
        if f["m5_trans"]:
            last = f["m5_trans"][-1]
            out.append(f"**Parametric driver:** the worst health-zone step ({last[1]}, {last[0]}) was driven by **{last[2]}** — the signal to watch on this truck.")
        if np.isfinite(f["cdelta"]) and f["cdelta"] <= -0.3:
            out.append(f"**Charging weakened:** regulation ceiling fell {abs(f['cdelta']):.1f} V early→late — physical corroboration of alternator wear.")
    else:
        if "watch" in f["arch"].lower():
            sig = f["breaches"][0][0] if f["breaches"] else "M5 health-zone"
            out.append(f"**Watch-list.** In-service but {sig} is elevated vs its own baseline — inspect ahead of the ~{f['median_rul']:.0f} d fleet estimate.")
        else:
            out.append(f"**Routine.** No actionable signal; follow the {f['fw_p25']:.0f}–{f['fw_p75']:.0f} d fleet window. RUL *band* (not the point estimate) is the trustworthy output.")
        if np.isfinite(f["cdelta"]) and f["cdelta"] <= -0.3:
            out.append(f"Charging ceiling has slipped {abs(f['cdelta']):.1f} V — early physical sign worth a recheck.")
    return out


# ---------------------------------------------------------------- how-to content
HOWTO = """\
## How to read these graphs (applies to every VIN)

Each figure stacks **five time-aligned panels** that share one calendar x-axis (Panel 4 uses engine RPM).
Vertical guide lines mark the same moments across panels: **grey dotted** = last data, **orange dashed** =
first physics deviation, **dark** = failure (red ✕) or forecast.

> **The one honest idea to hold onto:** Panel 1's RUL line is a **fleet clock** — a schedule based on how
> long alternators in this fleet last, *not* a per-truck prediction. The truck's *own* condition lives in
> Panels 1b–4. When the schedule and the physics disagree, **trust the physics.**

### Panel 1 — Remaining life: the FLEET SCHEDULE (age-based, not per-truck)
- **Simple:** how many days the *fleet average* says are left — a **schedule**, not a per-truck prediction.
  It slides down with age toward the failure ✕ (or a dotted forecast for healthy trucks); right axis = km.
- **Technical:** conditional Weibull posterior median + 80 % band over age; horizon zones at 180/90/30 d.
- **Read the over-prediction — don't be fooled by the drop.** For these premature failures the schedule still
  showed **130–270 days of life at the last reading**, so the curve drops sharply to the actual failure ✕,
  labelled *"OVER-PREDICTED ~X d expected vs Y d actual."* That gap is **honest and expected**: an age-based
  fleet clock cannot see an individual truck about to fail. It is **not** smoothed into a fake gradual decline.
  This is exactly why per-truck RUL is reported as a **band + the 601 d fleet window — never a per-truck
  promise** — and why the condition panels (1b–4) are where the real per-truck warning lives.

### Panel 1b — Health zone (parameter-driven M5)
- **Simple:** a single "health score" built from the truck's own voltage and excitation behaviour. When it
  steps up a zone, a numbered tag says **what moved it** (e.g. "→ORANGE · GED2 rate").
- **Technical:** M5 = weighted sum of VSI deviation, VSI volatility, GED2 rate, VSI range (recalibrated,
  thresholds 0.15/0.35/0.55). The transition driver is the largest-contributing component at that date.
  This is the parameter-decided zone view (the weaker system — it can flag early/transiently — so it is
  supporting evidence, not the primary alarm).

### Panel 2 — Voltage health (condition)
- **Simple:** the truck's charging voltage over time. The green band (27–29 V) is healthy; the red dashed
  line (24 V) is a sag threshold; tick marks at the bottom are days it sagged.
- **Technical:** daily VSI mean + p05–p95 envelope; 28.2 V DICV nominal; resting (engine-off) voltage
  dashed. Persistent drift below the band or rising sag density indicates regulation/charging loss.

### Panel 3 — Precursor signals (physics early-warning)
- **Simple:** several "is-it-getting-worse?" signals, all drawn so **up = worse**. A signal in **bold colour**
  has crossed the worry line; **grey** signals are quiet. If nothing crosses, it says *"No precursor detected."*
- **Technical:** each signal as **σ above the truck's own healthy baseline** (smoothed). The σ=2 dashed line
  is the discriminative gate; the legend lists breaching signals; the GED2 storm is marked separately. The
  orange callout gives the first-deviation lead time.

### Panel 4 — Charging signature (the "why")
- **Simple:** the alternator's voltage at different engine speeds, **early life (green)** vs **late life (red)**.
  If the red line drops below the green, the alternator is losing its ability to hold voltage.
- **Technical:** median VSI binned by RPM (600–2500). A late-life drop in the high-RPM plateau ("ceiling")
  is the physical fingerprint of stator/diode/regulation wear; overlapping lines = regulation still intact.
"""


def archetype_legend():
    return ("**Archetypes:** *Early-warning (GED2 storm)* — rare, gives short-horizon notice · "
            "*Gradual precursor* — a signal drifts over weeks/months · *Silent/abrupt* — no warning, the hard case · "
            "*In-service watch / quiet* — healthy trucks with vs without an elevated signal.")


# ---------------------------------------------------------------- build
def main():
    rows = [facts(v) for v in ALL]
    json.dump(rows, open(os.path.join(RESULTS, "V11.2_ALT_evidence_stack_facts.json"), "w"), indent=2, default=str)

    md = [f"# Alternator V11.2 — Per-VIN RUL + Physics Evidence Stack: Reading Guide & Insights",
          f"\n**Date:** {DATE}  ·  **Figures:** `V11.2_ALT/visualizations/rul_evidence_stack/`  ·  "
          f"**Fleet:** 25 trucks (10 failed VIN1–10, 15 in-service VIN11–25)\n", HOWTO,
          "\n## Fleet summary\n",
          "| VIN | Status | Risk | Archetype | Precursor (at failure / now) | M5 (final · peak) | Schedule over-pred | Charging Δ |",
          "|---|---|---|---|---|---|---|---|"]
    for f in rows:
        if f["breaches"]:
            fd = f"{f['breaches'][0][0]} " + (f"~{f['breaches'][0][1]}d" if f["failed"] else "(now)")
        elif f["ged"]:
            fd = f"GED2 ~{f['ged_lead']}d"
        else:
            fd = "—"
        cd = f"{f['cdelta']:+.2f} V" if np.isfinite(f["cdelta"]) else "n/a"
        m5cell = f["m5_final"] + (f" · peak {f['m5_peak']}" if f.get("m5_peak") else "")
        op = f"~{f['over_pred']:.0f} d" if f.get("over_pred") is not None else "—"
        md.append(f"| {f['dname']} | {'FAILED' if f['failed'] else 'in-service'} | {f['risk_band']} | "
                  f"{f['arch']} | {fd} | {m5cell} | {op} | {cd} |")
    md.append("\n" + archetype_legend() + "\n")
    md.append("\n## Per-VIN reading & insights\n")
    for f in rows:
        md.append(f"\n### {f['dname']}  ·  _{f['arch']}_\n")
        md.append(f"![{f['dname']} evidence stack]({VIZ_REL}/{f['dname']}_evidence_stack.png)\n")
        md.append(f"**In simple terms.** {author_simple(f)}\n")
        md.append(f"**Technically.** {author_technical(f)}\n")
        md.append("**Key insights.**")
        for ins in author_insights(f):
            md.append(f"- {ins}")
        md.append("")
        snap = f["snap"]
        md.append("**Snapshot — every component at each key moment (read down a column):**\n")
        md.append("| Component | " + " | ".join(f"{lab} ({dt.date()})" for lab, dt in snap["cols"]) + " |")
        md.append("|---|" + "---|" * len(snap["cols"]))
        for rname, cells in snap["rows"].items():
            md.append(f"| {rname} | " + " | ".join(c[0] for c in cells) + " |")
        md.append("")
    md.append("\n## Honest limitations\n"
              "- Per-truck RUL point estimates do not beat the fleet clock (MAE ~140 d vs ~50 d); the **band** "
              "and the **601 d / ~120k km window** are the timing deliverables.\n"
              "- Precursors fire for a minority of failures (GED2 2/10; compound 3/10); most failures are abrupt.\n"
              "- The M5 health-zone is supporting context only (weak sensitivity, can flag early/transiently).\n"
              "- All signals carry the n=25 / 10-event data ceiling; figures are decision-support, not guarantees.\n")
    md_text = "\n".join(md)
    md_path = os.path.join(REPORTS, "V11.2_ALT_evidence_stack_reading_guide.md")
    open(md_path, "w", encoding="utf-8").write(md_text)
    print("wrote", md_path, len(md_text.split()), "words")

    # ---- DOCX ----
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        doc = Document()
        doc.add_heading("Alternator V11.2 — Per-VIN RUL + Physics Evidence Stack", 0)
        doc.add_paragraph(f"Reading Guide & Insights  ·  {DATE}  ·  25 trucks (failed VIN1–10, in-service VIN11–25)")
        # how-to (strip markdown markers lightly)
        for block in HOWTO.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].split("\n")[0], level=2)
                rest = "\n".join(block.split("\n")[1:]).strip()
                if rest:
                    doc.add_paragraph(rest.replace("**", "").replace("- ", "• "))
            elif block.startswith("## "):
                doc.add_heading(block[3:], level=1)
            elif block.startswith(">"):
                p = doc.add_paragraph(block.replace(">", "").replace("**", "").strip()); p.italic = True
            else:
                doc.add_paragraph(block.replace("**", "").replace("- ", "• "))
        # fleet table
        doc.add_heading("Fleet summary", level=1)
        cols = ["VIN", "Status", "Risk", "Archetype", "Precursor", "M5 (final·peak)", "Sched over-pred", "Charging Δ"]
        t = doc.add_table(rows=1, cols=len(cols)); t.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            t.rows[0].cells[i].text = c
        for f in rows:
            if f["breaches"]:
                fd = f"{f['breaches'][0][0]} " + (f"~{f['breaches'][0][1]}d" if f["failed"] else "(now)")
            elif f["ged"]:
                fd = f"GED2 ~{f['ged_lead']}d"
            else:
                fd = "—"
            cd = f"{f['cdelta']:+.2f} V" if np.isfinite(f["cdelta"]) else "n/a"
            m5cell = f["m5_final"] + (f" · peak {f['m5_peak']}" if f.get("m5_peak") else "")
            op = f"~{f['over_pred']:.0f} d" if f.get("over_pred") is not None else "—"
            cells = t.add_row().cells
            for i, v in enumerate([f["dname"], "FAILED" if f["failed"] else "in-service", f["risk_band"],
                                   f["arch"], fd, m5cell, op, cd]):
                cells[i].text = str(v)
        # per-VIN
        doc.add_heading("Per-VIN reading & insights", level=1)
        for f in rows:
            doc.add_page_break()
            doc.add_heading(f"{f['dname']} — {f['arch']}", level=2)
            img = os.path.join(VIZ_ABS, f"{f['dname']}_evidence_stack.png")
            if os.path.exists(img):
                doc.add_picture(img, width=Inches(5.2))
            doc.add_paragraph().add_run("In simple terms. ").bold = True
            doc.paragraphs[-1].add_run(author_simple(f).replace("**", ""))
            doc.add_paragraph().add_run("Technically. ").bold = True
            doc.paragraphs[-1].add_run(author_technical(f).replace("**", ""))
            doc.add_paragraph().add_run("Key insights.").bold = True
            for ins in author_insights(f):
                doc.add_paragraph(ins.replace("**", ""), style="List Bullet")
            snap = f["snap"]
            doc.add_paragraph().add_run("Snapshot — every component at each key moment:").bold = True
            st = doc.add_table(rows=1, cols=len(snap["cols"]) + 1); st.style = "Light Grid Accent 1"
            st.rows[0].cells[0].text = "Component"
            for j, (lab, dt) in enumerate(snap["cols"], start=1):
                st.rows[0].cells[j].text = f"{lab}\n{dt.date()}"
            for rname, cells in snap["rows"].items():
                rc = st.add_row().cells; rc[0].text = rname
                for j, c in enumerate(cells, start=1):
                    rc[j].text = c[0]
        docx_path = os.path.join(REPORTS, "V11.2_ALT_evidence_stack_reading_guide.docx")
        doc.save(docx_path)
        print("wrote", docx_path)
    except Exception as ex:
        print("DOCX step skipped:", ex)


if __name__ == "__main__":
    main()
